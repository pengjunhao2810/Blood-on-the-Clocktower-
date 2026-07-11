import sys, os, json
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
# 添加 werewolf_ai 路径以便导入 games.blood_on_clocktower 包
_werewolf_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '本地训练的ai', 'werewolf_ai'))
if os.path.isdir(_werewolf_path):
    sys.path.insert(0, _werewolf_path)
from flask import Flask, jsonify, request, render_template_string, make_response, send_file
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime
from games.blood_on_clocktower.rules import BloodOnClocktowerGame
from games.blood_on_clocktower.roles import BOTC_ROLES, BOTC_TEAMS
from core.agent import SocialDeductionAgent

app = Flask(__name__)

CUSTOM_NAMES = ["任","mh","Youmi","则夷则","欧拉函数","Hannoy","Star","丘皮卡","戒色大师","猪头鹰","progress","wowo","大吴吴","Run","good name","Rujie Shi","大美","琥珀王","云面包","kiod","薛猫","大6","大太阳","困困","瑜瑜","when？","香辣鸡腿堡","Takai&xingfei","澳大利亚","吃鱼的猫"]

class GameController:
    def __init__(self, num_players=7):
        self.num_players = num_players
        self.agents = None
        self.finished = False
        self._step_count = 0
        self._init_game()

    def _init_game(self):
        import random as _r
        shuffled = list(CUSTOM_NAMES)
        _r.shuffle(shuffled)
        names = shuffled[:self.num_players]
        self.agents = [SocialDeductionAgent(name, '?') for name in names]
        self.game = BloodOnClocktowerGame(num_players=self.num_players, script='暗流涌动')
        self.game.setup_game(self.agents)
        self.finished = False
        self.flow = ['role_assign']
        self.day_prepared = False
        if not hasattr(self, 'full_log'):
            self.full_log = []
        self.full_log.append({'num_players': self.num_players, 'log': list(self.game.storyteller_log)})
        while len(self.full_log) > 5:
            self.full_log.pop(0)

    def step(self):
        self._step_count += 1
        if self.finished or self.game.game_record.get('result'):
            self.finished = True
            return self.get_state()

        if not self.flow:
            self.flow = ['night', 'private_chat', 'public_chat', 'nomination']

        try:
            action = self.flow.pop(0)

            if action == 'role_assign':
                self.game.phase_role_assignment()

            elif action == 'night':
                if self.game.day_count > 0:
                    self.game.end_day()
                    if self.game.game_record.get('result'):
                        self.finished = True
                        return self.get_state()
                self.game.run_night()
                if self.game.game_record.get('result'):
                    self.finished = True

            elif action == 'private_chat':
                if not self.day_prepared:
                    self.game.start_day()
                    self.day_prepared = True
                    if self.game.game_record.get('result'):
                        self.finished = True
                        return self.get_state()
                self.game._private_chat_phase()

            elif action == 'public_chat':
                self.game._public_chat_phase()

            elif action == 'nomination':
                self.game._nomination_and_voting_phase()
                self.day_prepared = False
        except Exception as e:
            import traceback
            with open(os.path.join(os.path.dirname(__file__), 'crash.log'), 'a', encoding='utf-8') as _f:
                _f.write(f"=== step() exception ===\n")
                traceback.print_exc(file=_f)
            self.finished = True
            if not self.game.game_record.get('result'):
                self.game.game_record['result'] = 'error'

        return self.get_state()

    def get_state(self):
        g = self.game
        team_cn = {'townsfolk': '镇民', 'outsider': '外来者', 'minion': '爪牙', 'demon': '恶魔'}
        phase_icons = {
            'SETUP': '⚙️', 'ROLE_ASSIGN': '📋',
            'NIGHT': '🌙', 'DAY': '☀️',
            'PRIVATE_CHAT': '🤫', 'PUBLIC_CHAT': '💬',
            'NOMINATION': '🗳️',
        }

        raw = g.game_phase
        base = raw.split('_')[0] if '_' in raw else raw
        icon = phase_icons.get(base, '📜')

        phase_labels = {
            'SETUP': ('⚙️ 游戏设置', '准备中...'),
            'ROLE_ASSIGN': ('📋 角色分配', '各玩家已查看身份'),
        }
        if base in phase_labels:
            pn, pd = phase_labels[base]
        elif 'NIGHT' in raw:
            pn, pd = '🌙 天黑闭眼', '各角色按顺序行动中...'
        elif 'PRIVATE_CHAT' in raw:
            pn, pd = '🤫 私聊环节', '玩家私下交流信息'
        elif 'PUBLIC_CHAT' in raw:
            pn, pd = '💬 公聊环节', '玩家公开讨论'
        elif 'NOMINATION' in raw:
            pn, pd = '🗳️ 提名投票', '提名·辩护·全体表决'
        elif 'DAY' in raw:
            pn, pd = '☀️ 白天', '交流与投票'
        else:
            pn, pd = raw, ''

        alive = g.get_alive_names()
        all_agents = g.registry.all_agents()
        all_names = [a.name for a in all_agents]
        evil_agents = [a for a in all_agents if a.role in BOTC_TEAMS['demon'] + BOTC_TEAMS['minion']]

        players = []
        seat_order = g.player_order
        for a in all_agents:
            info = BOTC_ROLES.get(a.role, {})
            seat = (seat_order.index(a.name) + 1) if a.name in seat_order else 0
            drunk_info = ''
            if a.role == "酒鬼" and a.alive:
                fake_role = a.game_state.get("fake_role", "")
                if fake_role:
                    fake_icon = BOTC_ROLES.get(fake_role, {}).get('icon', '?')
                    drunk_info = f"{fake_role} {fake_icon}"
            players.append({
                'name': a.name, 'seat': seat, 'role': a.role, 'icon': info.get('icon', '?'),
                'team': info.get('team', ''),
                'team_cn': team_cn.get(info.get('team', ''), ''),
                'status': 'alive' if a.alive else 'dead',
                'poisoned': a.game_state.get('is_poisoned', False) if a.alive else False,
                'drunk_info': drunk_info,
            })

        # 私聊历史整理成前端友好的格式
        chat_threads = []
        pch = g.game_record.get("private_chat_history", {})
        for thread_key, msgs in pch.items():
            # thread_key格式: D1_玩家1🔄玩家3
            parts = thread_key.split("_", 1)
            day_label = parts[0] if len(parts) > 1 else "D0"
            pair_label = parts[1] if len(parts) > 1 else thread_key
            chat_threads.append({
                "key": thread_key,
                "day": day_label.replace("D", "第") + "天",
                "pair": pair_label,
                "messages": msgs,
            })

        # 酒鬼列表
        drunk_players = []
        for a in all_agents:
            if a.role == "酒鬼" and a.alive:
                fake_role = a.game_state.get("fake_role", "")
                if fake_role:
                    drunk_players.append({
                        'name': a.name,
                        'fake_role': fake_role,
                        'fake_icon': BOTC_ROLES.get(fake_role, {}).get('icon', '?'),
                    })
        # 干扰项（占卜师）
        red_herring = ""
        for a in all_agents:
            if a.role == "占卜师" and a.alive:
                rh_name = a.game_state.get("red_herring", "")
                if rh_name:
                    red_herring = f"{rh_name}"
                break

        role_configs = {
            6: {"townsfolk": 3, "outsider": 1, "minion": 1, "demon": 1},
            7: {"townsfolk": 5, "outsider": 0, "minion": 1, "demon": 1},
            8: {"townsfolk": 5, "outsider": 0, "minion": 2, "demon": 1},
            9: {"townsfolk": 5, "outsider": 2, "minion": 1, "demon": 1},
            10: {"townsfolk": 7, "outsider": 0, "minion": 2, "demon": 1},
            11: {"townsfolk": 7, "outsider": 1, "minion": 2, "demon": 1},
            12: {"townsfolk": 7, "outsider": 2, "minion": 2, "demon": 1},
        }

        # 构建公开日志（过滤敏感夜间信息，仅保留公共可见内容）
        raw_log = g.storyteller_log[-200:] if g.storyteller_log else []
        filtered_log = []
        import re as _re
        _public_keep = [
            _re.compile(r'={10,}'),             # 分隔线（可能前面有\n）
            _re.compile(r'昨晚死亡|平安夜'),     # 死亡公告
            _re.compile(r'公聊环节'),             # 公聊阶段标记
            _re.compile(r'私聊环节'),             # 私聊阶段标记
            _re.compile(r'^  [^\[(]+:'),        # 公聊发言（玩家名: 内容，排除角色分配行）
            _re.compile(r'\[投票\]|\[计票\]'),         # 投票信息（不含投票详情，含阵营信息）
            _re.compile(r'被处决|未被处决'),     # 处决结果
            _re.compile(r'提名了|\[提名发言\]|\[辩护发言\]|发起提名'),  # 提名/辩护发言
            _re.compile(r'\[猎手\]'),            # 猎手开枪（公开事件）
            _re.compile(r'\[小丑\]|\[圣徒\]'),    # 特殊角色公开触发
            _re.compile(r'\[设置\]|\[ML训练\]'), # 设置信息
            _re.compile(r'游戏结束|胜利|落败'),   # 游戏结果
            _re.compile(r'给各个玩家发送信息'),   # 首夜信息简明提示
            _re.compile(r'--- \['),              # 阶段分隔符（可能前面有\n）
        ]
        _sensitive_lines = [
            _re.compile(r'\(恶魔\)$'),   # 角色分配行: "  玩家名: 角色 (恶魔)"
            _re.compile(r'\(爪牙\)$'),
            _re.compile(r'\(镇民\)$'),
            _re.compile(r'\(外来者\)$'),
            _re.compile(r'\(间谍\)查看到全角色'),
            _re.compile(r'\(恶魔\)已知爪牙'),
            _re.compile(r'可伪装角色'),
            _re.compile(r'投票详情:'),                 # 投票详情含玩家阵营
        ]
        for line in raw_log:
            if any(p.search(line) for p in _sensitive_lines):
                continue
            if any(p.search(line) for p in _public_keep):
                filtered_log.append(line)

        return {
            'phase_name': pn, 'phase_desc': pd, 'raw_phase': raw,
            'day_count': g.day_count, 'night_count': g.night_count,
            'alive_count': len(alive), 'dead_count': len(all_names) - len(alive),
            'evil_count': len([a for a in evil_agents if a.alive]),
            'players': players,
            'log': filtered_log,
            'finished': self.finished or bool(g.game_record.get('result')),
            'result': g.game_record.get('result'),
            'chat_threads': chat_threads,
            'drunk_players': drunk_players,
            'red_herring': red_herring,
            'role_config': role_configs,
            'current_player_count': self.num_players,
            'step_count': getattr(self, '_step_count', 0),
            'nomination_history': g.game_record.get("nomination_history", {}),
            'vote_history': g.game_record.get("vote_history", {}),
        }

ctrl = GameController()

HTML = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>血染钟楼 · 魔典</title>
<link href="https://fonts.googleapis.com/css?family=Roboto+Condensed&display=swap" rel="stylesheet">
<style>
* { margin:0; padding:0; box-sizing:border-box; }
body { font-family:'Roboto Condensed','Microsoft YaHei',sans-serif; background:#0a0a0f; color:#d0d0d0; min-height:100vh; overflow-x:hidden; }
.header { background:linear-gradient(180deg,#1a0000,#0a0a0f); padding:14px 30px; display:flex; justify-content:space-between; align-items:center; border-bottom:2px solid #cc000044; }
.header h1 { color:#cc0000; font-size:22px; letter-spacing:3px; text-transform:uppercase; font-weight:300; }
.header h1 span { color:#fff; font-weight:700; }
.header .info { color:#888; font-size:12px; }
.header .info span { color:#cc0000; }
.container { max-width:1200px; margin:0 auto; padding:16px; }
.phase-bar { background:linear-gradient(135deg,#1a0000,#0a0000); border:1px solid #cc000044; border-radius:8px; padding:10px 20px; margin-bottom:12px; display:flex; justify-content:space-between; align-items:center; }
.phase-bar .phase-name { color:#cc0000; font-size:18px; font-weight:bold; text-transform:uppercase; letter-spacing:2px; }
.phase-bar .phase-desc { color:#888; font-size:12px; margin-top:2px; }
.phase-bar .round { color:#666; font-size:13px; }
.phase-bar .round span { color:#cc0000; font-weight:bold; }
.controls { display:flex; gap:8px; margin-bottom:12px; flex-wrap:wrap; }
.controls button { padding:10px 24px; border:none; border-radius:4px; font-size:14px; font-weight:bold; cursor:pointer; transition:all .2s; font-family:'Roboto Condensed',sans-serif; text-transform:uppercase; letter-spacing:1px; }
.controls button:hover { transform:translateY(-1px); box-shadow:0 2px 10px rgba(0,0,0,.5); }
.btn-step { background:#cc0000; color:#fff; }
.btn-step:disabled { opacity:0.4; cursor:not-allowed; transform:none !important; }
.btn-auto { background:#333; color:#ccc; }
.btn-reset { background:#cc0000; color:#fff; }
.btn-export { background:#333; color:#ccc; }
.info-row { display:flex; gap:16px; margin-bottom:12px; flex-wrap:wrap; }
.info-item { background:#111; border:1px solid #cc000022; border-radius:4px; padding:8px 16px; color:#888; font-size:12px; }
.info-item span { color:#cc0000; font-weight:bold; }
/* ============ 玩家圆形排列(魔典风格) ============ */
.circle-container { position:relative; width:100%; max-width:700px; margin:0 auto; aspect-ratio:1/1; }
.player-card { position:absolute; width:90px; height:90px; transform:translate(-50%,-50%); cursor:pointer; transition:all .3s; z-index:1; }
.player-card:hover { z-index:10; transform:translate(-50%,-50%) scale(1.3); }
.player-card.zoomed { transform:translate(-50%,-50%) scale(2.2); z-index:100; }
.player-card.zoomed ~.player-card { opacity:0.1; }
.player-card .circle-bg { width:100%; height:100%; border-radius:50%; background:linear-gradient(135deg,#111,#0a0a0a); border:3px solid #333; display:flex; flex-direction:column; align-items:center; justify-content:center; text-align:center; overflow:hidden; padding:4px; transition:all .3s; }
.player-card.alive .circle-bg { border-color:#444; box-shadow:0 0 6px rgba(204,0,0,.15); }
.player-card.dead .circle-bg { border-color:#333; opacity:.35; }
.player-card.dead .circle-bg::after { content:'✕'; position:absolute; font-size:40px; color:#cc0000; opacity:.8; }
.player-card .icon { font-size:22px; line-height:1; }
.player-card .name { font-size:10px; font-weight:bold; margin:1px 0 0; line-height:1.1; }
.player-card .role { font-size:8px; padding:1px 5px; border-radius:4px; display:inline-block; margin-top:1px; white-space:nowrap; }
.player-card .role.townsfolk { background:#1a3a5c; color:#5c9ce0; }
.player-card .role.outsider { background:#3a1a5c; color:#9c5ce0; }
.player-card .role.minion { background:#5c1a1a; color:#e05c5c; }
.player-card .role.demon { background:#5c0000; color:#ff4444; }
.player-card .status { font-size:8px; padding:0 4px; border-radius:3px; margin-top:1px; }
.player-card .status.alive { color:#444; }
.player-card .status.dead { color:#cc0000; }
.player-card .poisoned { position:absolute; top:-3px; right:-3px; background:#cc0000; color:#fff; font-size:8px; width:18px; height:18px; border-radius:50%; display:flex; align-items:center; justify-content:center; z-index:2; }

/* ============ 双栏布局: 日志 | 玩家圆环 ============ */
.main-layout { display:grid; grid-template-columns:1.1fr 1fr; gap:14px; align-items:start; min-height:calc(100vh - 200px); }

/* ============ 游戏设置信息(酒鬼+干扰项) ============ */
.setup-info { background:linear-gradient(135deg,#1a1a2e,#0d0d1a); border:1px solid #2a2a4e; border-radius:8px; padding:8px 14px; margin-bottom:10px; font-size:12px; color:#aaa; visibility:hidden; min-height:44px; }
.setup-info.show { visibility:visible; }
.setup-info .si-row { display:flex; align-items:center; gap:8px; padding:2px 0; }
.setup-info .si-icon { font-size:14px; flex-shrink:0; }
.setup-info .si-name { color:#e0e0e0; font-weight:bold; }
.setup-info .si-arrow { color:#888; }
.setup-info .si-role { color:#7ab8ff; }
.setup-info .si-truth { color:#ff7a7a; }

/* ============ 日志(左) - 魔典暗色 ============ */
.log-box { background:#050508; border:1px solid #1a1a1a; border-radius:4px; display:flex; flex-direction:column; max-height:calc(100vh - 140px); padding:6px 8px; }
.log-box .log-title { color:#666; font-size:11px; margin-bottom:3px; padding-bottom:3px; border-bottom:1px solid #1a1a1a; display:flex; align-items:center; justify-content:space-between; flex-shrink:0; }
.log-box .log-title .log-controls label { font-size:10px; color:#555; cursor:pointer; display:flex; align-items:center; gap:2px; user-select:none; }
.log-box .log-title .log-controls input[type=checkbox] { accent-color:#cc0000; cursor:pointer; }
.log-filter { display:flex; gap:2px; margin-bottom:3px; flex-wrap:wrap; flex-shrink:0; }
.log-filter button { background:#111; border:1px solid #222; color:#666; font-size:9px; padding:1px 6px; border-radius:3px; cursor:pointer; }
.log-filter button:hover { background:#1a1a1a; color:#aaa; }
.log-filter button.active { background:#cc000022; border-color:#cc000044; color:#cc0000; }
.log-box .log-entries { overflow-y:auto; flex:1; }
.log-box .log-entry { padding:1px 5px; font-size:11px; line-height:1.4; border-radius:1px; }
.log-box .log-entry.hidden { display:none; }
.log-box .log-entry.entry-phase { color:#cc0000; font-weight:bold; background:rgba(204,0,0,.06); border-left:2px solid #cc0000; }
.log-box .log-entry.entry-night { color:#888; background:rgba(255,255,255,.04); border-left:2px solid #555; }
.log-box .log-entry.entry-day { color:#ccc; background:rgba(255,255,255,.03); border-left:2px solid #666; }
.log-box .log-entry.entry-death { color:#ff4444; background:rgba(255,0,0,.06); border-left:2px solid #cc0000; }
.log-box .log-entry.entry-chat { color:#aaa; background:rgba(255,255,255,.02); border-left:2px solid #444; }
.log-box .log-entry.entry-vote { color:#cc8833; background:rgba(204,136,51,.06); border-left:2px solid #cc8833; }
.log-box .log-entry.entry-action { color:#888; background:rgba(255,255,255,.03); border-left:2px solid #555; }
.log-box .log-entry.entry-result { color:#cc0000; background:rgba(204,0,0,.1); border-left:2px solid #cc0000; font-weight:bold; }
.log-box .log-entry.entry-system { color:#555; }

.right-panel { min-width:0; }

.overlay { display:none; position:fixed; top:0; left:0; right:0; bottom:0; background:rgba(0,0,0,.9); z-index:1000; justify-content:center; align-items:center; }
.overlay.show { display:flex; }
.overlay .result-box { background:#0a0a0a; border:2px solid #cc0000; border-radius:4px; padding:40px; text-align:center; max-width:450px; }
.overlay .result-box h2 { font-size:28px; margin-bottom:12px; text-transform:uppercase; letter-spacing:2px; }
.overlay .result-box .good { color:#ccc; }
.overlay .result-box .evil { color:#cc0000; }
.overlay .result-box p { color:#888; font-size:14px; margin-bottom:20px; }
.overlay .result-box button { padding:12px 40px; border:none; border-radius:4px; background:#cc0000; color:#fff; font-weight:bold; font-size:14px; cursor:pointer; text-transform:uppercase; letter-spacing:1px; }

@keyframes pulse { 0%,100%{opacity:1} 50%{opacity:.5} }
.playing .phase-name { animation:pulse 1.5s infinite; }

/* ============ 角色配置栏 ============ */
.role-config-bar { background:#111125; border:1px solid #2a2a4e; border-radius:8px; margin-bottom:12px; overflow:hidden; }
.role-config-bar .rc-title { padding:8px 14px; cursor:pointer; display:flex; justify-content:space-between; align-items:center; font-size:13px; color:#aaa; user-select:none; transition:background .15s; }
.role-config-bar .rc-title:hover { background:#1a1a3e; }
.role-config-bar .rc-title span:first-child { color:#ccc; }
.role-config-bar .rc-arrow { font-size:11px; color:#666; transition:transform .2s; }
.role-config-bar .rc-arrow.open { transform:rotate(90deg); }
.role-config-bar .rc-table { display:none; padding:0 14px 10px; }
.role-config-bar .rc-table.show { display:block; }
.role-config-bar .rc-row { display:grid; grid-template-columns:70px repeat(4,1fr); gap:4px; padding:3px 0; font-size:12px; color:#888; text-align:center; }
.role-config-bar .rc-header { color:#ffd700; font-weight:bold; border-bottom:1px solid #2a2a4e; padding-bottom:5px; margin-bottom:2px; }
.role-config-bar .rc-row .tf { color:#7ab8ff; }
.role-config-bar .rc-row .os { color:#a89bff; }
.role-config-bar .rc-row .mi { color:#ff7a7a; }
.role-config-bar .rc-row .de { color:#ff8888; }
.role-config-bar .rc-row.rc-current { background:#1a1a3e; border-radius:4px; font-weight:bold; color:#e0e0e0; }

/* ============ 滚动条 ============ */
::-webkit-scrollbar { width:6px; }
::-webkit-scrollbar-track { background:transparent; }
::-webkit-scrollbar-thumb { background:#333; border-radius:3px; }
::-webkit-scrollbar-thumb:hover { background:#555; }

/* ============ 响应式: 窄屏折叠为单栏 ============ */
@media (max-width:900px) {
    .main-layout { grid-template-columns:1fr; }
    .right-panel { order:-1; }
    .right-panel .circle-container { max-width:400px; }
    .controls { justify-content:center; }
    .controls button { padding:10px 18px; font-size:13px; }
    .header { flex-direction:column; gap:10px; text-align:center; padding:12px 16px; }
    .header h1 { font-size:18px; }

}
</style>
</head>
<body>
<div class="header">
    <h1>🕯️ 血染钟楼</h1>
    <div class="info" style="display:flex;align-items:center;gap:10px;">
        <select id="playerSelect" onchange="setPlayerCount(this.value)" style="background:#2a2a5e;color:#ffd700;border:2px solid #ffd700;border-radius:6px;padding:6px 12px;font-size:16px;font-weight:bold;cursor:pointer;">
            <option value="5">5人局</option>
            <option value="6">6人局</option>
            <option value="7" selected>7人局</option>
            <option value="8">8人局</option>
            <option value="9">9人局</option>
            <option value="10">10人局</option>
            <option value="11">11人局</option>
            <option value="12">12人局</option>
        </select>
        <a href="/private_chat" target="_blank" style="background:#07c160;color:#fff;text-decoration:none;padding:6px 14px;border-radius:6px;font-size:13px;font-weight:bold;white-space:nowrap;">💬 私聊</a>
    </div>
</div>
<div class="container">
    <div id="phaseBar" class="phase-bar">
        <div>
            <div id="phaseName" class="phase-name">📋 准备开始</div>
            <div id="phaseDesc" class="phase-desc">点击下方「下一步」逐帧推进游戏</div>
        </div>
        <div class="round" id="roundInfo">🌙 第 <span id="nightCount">0</span> 晚 · ☀️ 第 <span id="dayCount">0</span> 天 · ⏱ <span id="stepCount">0</span> 步</div>
    </div>
    <div class="role-config-bar" id="roleConfigBar">
        <div class="rc-title" onclick="toggleRoleConfig()">
            <span id="rcSummary">📋 镇民 5 · 外来者 0 · 爪牙 1 · 恶魔 1</span>
            <span class="rc-arrow" id="rcArrow">▶</span>
        </div>
        <div class="rc-table" id="rcTable">
            <div class="rc-row rc-header">
                <span>玩家数</span><span>镇民</span><span>外来者</span><span>爪牙</span><span>恶魔</span>
            </div>
        </div>
    </div>
    <div class="setup-info" id="setupInfo"></div>
    <div class="info-row">
        <div class="info-item">🧑‍🤝‍🧑 存活: <span id="aliveCount">7</span>人</div>
        <div class="info-item">💀 死亡: <span id="deadCount">0</span>人</div>
        <div class="info-item">👿 邪恶存活: <span id="evilCount">-</span>人</div>
    </div>
    <div class="controls">
        <button id="btnStep" class="btn-step" onclick="stepGame()">▶ 下一步</button>
        <button id="btnAuto" class="btn-auto" onclick="toggleAuto()">▶ 自动播放</button>
        <button class="btn-reset" onclick="resetGame()">↻ 新游戏</button>
        <button class="btn-export" onclick="exportDoc()">📄 导出Word</button>
    </div>

    <!-- 主体布局: 日志 | 私聊 | 玩家圆环 -->
    <div class="main-layout">
        <div class="log-box" id="logArea">
            <div class="log-title">
                <span>📜 日志</span>
                <span class="log-controls">
                    <label><input type="checkbox" id="autoScrollToggle" checked onchange="toggleAutoScroll()"> 自动滚动</label>
                </span>
            </div>
            <div class="log-filter" id="logFilter">
                <button class="active" data-filter="all" onclick="filterLog('all',this)">全部</button>
                <button data-filter="phase" onclick="filterLog('phase',this)">📋</button>
                <button data-filter="action" onclick="filterLog('action',this)">⚡</button>
                <button data-filter="chat" onclick="filterLog('chat',this)">💬</button>
                <button data-filter="vote" onclick="filterLog('vote',this)">🗳️</button>
                <button data-filter="death" onclick="filterLog('death',this)">💀</button>
            </div>
            <div class="log-entries" id="logEntries"></div>
        </div>
        <div class="right-panel">
            <div id="circleContainer" class="circle-container"></div>
        </div>
    </div>
</div>
<div id="overlay" class="overlay">
    <div class="result-box">
        <h2 id="resultTitle" class="good">🏆 游戏结束</h2>
        <p id="resultDetail">善良阵营获胜！</p>
        <button onclick="closeOverlay()">确定</button>
    </div>
</div>
<script>
let autoPlaying=false, autoTimer=null, _stepping=false;
function updateBoard(d) {
    document.getElementById('phaseName').textContent = d.phase_name;
    document.getElementById('phaseDesc').textContent = d.phase_desc;
    document.getElementById('dayCount').textContent = d.day_count;
    document.getElementById('nightCount').textContent = d.night_count;
    document.getElementById('aliveCount').textContent = d.alive_count;
    document.getElementById('deadCount').textContent = d.dead_count;
    document.getElementById('evilCount').textContent = d.evil_count;

    renderSetupInfo(d);

    // 圆形排列玩家(最右侧) — 自适应间距不重叠
    let container = document.getElementById('circleContainer');
    let n = d.players.length;
    let cardSize = Math.max(68, Math.min(80, 180 - n * 8));
    let minGap = cardSize + 12;
    let idealR = minGap / (2 * Math.sin(Math.PI / n));
    let maxR = Math.min(container.offsetWidth / 2.3, container.offsetHeight / 2.3, 240);
    let radius = Math.min(idealR, maxR);
    let cx = container.offsetWidth / 2, cy = container.offsetHeight / 2;
    if(cy < 10) cy = 180;
    let h = '';
    for(let i = 0; i < n; i++) {
        let p = d.players[i];
        let angle = (i / n) * 2 * Math.PI - Math.PI / 2;
        let x = cx + radius * Math.cos(angle);
        let y = cy + radius * Math.sin(angle);
        let drunkIcon = p.drunk_info ? '<span class="icon-belief">'+p.drunk_info.split(' ').slice(1).join(' ')+'</span>' : '';
        h += '<div class="player-card '+p.status+'" style="left:'+x+'px;top:'+y+'px;width:'+cardSize+'px;height:'+cardSize+'px" onclick="toggleZoom(this)" title="点击放大">' +
            '<div class="circle-bg" style="font-size:'+(cardSize/12)+'px">' +
            '<div class="icon" style="font-size:'+Math.min(16,cardSize/4.5)+'px">'+p.icon+drunkIcon+'</div>' +
            '<div class="role '+p.team+'" style="font-size:'+Math.max(8,cardSize/11)+'px">'+p.role+(p.drunk_info?'🤔':'')+'</div>' +
            '<div class="seat" style="font-size:'+Math.max(11,cardSize/7)+'px;color:#4CAF50;font-weight:bold;">座位'+p.seat+'</div>' +
            '<div class="name" style="font-size:'+Math.max(9,cardSize/9)+'px">'+p.name+'</div>' +
            (p.poisoned?'<div class="poisoned">☠️</div>':'') +
            '</div></div>';
    }
    container.innerHTML = h;

    // === 说书人日志 ===
    window._fullLogData = d.log;
    renderLog(d.log);

    let btn = document.getElementById('btnStep');
    btn.disabled = d.finished;
    btn.textContent = d.finished ? '✓ 游戏结束' : '▶ 下一步';

    // 步数计数
    let sc = document.getElementById('stepCount');
    if(sc) sc.textContent = d.step_count || 0;

    updateRoleConfig(d);

    if(d.result) {
        setTimeout(() => {
            document.getElementById('overlay').classList.add('show');
            let t = document.getElementById('resultTitle'), dt = document.getElementById('resultDetail');
            if(d.result=='good_win') { t.textContent='🏆 善良阵营获胜！'; t.className='good'; dt.textContent='所有恶魔已被消灭！'; }
            else { t.textContent='💀 邪恶阵营获胜！'; t.className='evil'; dt.textContent='场上只剩两名存活玩家。'; }
        }, 600);
    }
}

function renderSetupInfo(d) {
    let el = document.getElementById('setupInfo');
    let parts = [];
    // 酒鬼
    if(d.drunk_players && d.drunk_players.length > 0) {
        for(let dp of d.drunk_players) {
            parts.push('<div class="si-row"><span class="si-icon">🍺</span><span class="si-name">'+dp.name+'</span><span class="si-arrow">认为自己</span><span class="si-role">'+dp.fake_role+' '+dp.fake_icon+'</span><span class="si-arrow">但其实是</span><span class="si-truth">酒鬼</span></div>');
        }
    }
    // 干扰项
    if(d.red_herring) {
        let rhPlayers = d.players.filter(p => p.name === d.red_herring);
        let rhRole = rhPlayers.length > 0 ? rhPlayers[0].role : '?';
        parts.push('<div class="si-row"><span class="si-icon">🎯</span><span class="si-name">'+d.red_herring+'</span><span class="si-arrow">是</span><span class="si-role">占卜师</span><span class="si-arrow">的干扰项（红鲱鱼）</span></div>');
    }
    if(parts.length > 0) {
        el.innerHTML = parts.join('');
        el.classList.add('show');
    } else {
        el.classList.remove('show');
    }
}

function toggleZoom(el) {
    let zoomed = el.classList.contains('zoomed');
    document.querySelectorAll('.player-card.zoomed').forEach(c => c.classList.remove('zoomed'));
    if(!zoomed) el.classList.add('zoomed');
}

function _escapeHtml(s) {
    if(!s) return '';
    return s.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');
}

function toggleRoleConfig() {
    let tbl = document.getElementById('rcTable');
    let arr = document.getElementById('rcArrow');
    tbl.classList.toggle('show');
    arr.classList.toggle('open');
}
function updateRoleConfig(d) {
    let cfg = d.role_config;
    let cur = d.current_player_count;
    if(!cfg || !cfg[cur]) return;
    let c = cfg[cur];
    let teamCn = {townsfolk:'镇民',outsider:'外来者',minion:'爪牙',demon:'恶魔'};
    let teamCls = {townsfolk:'tf',outsider:'os',minion:'mi',demon:'de'};
    document.getElementById('rcSummary').textContent =
        '\U0001F4CB 镇民 ' + c.townsfolk + ' \u00B7 外来者 ' + c.outsider + ' \u00B7 爪牙 ' + c.minion + ' \u00B7 恶魔 ' + c.demon;
    let rows = '<div class="rc-row rc-header"><span>玩家数</span><span>镇民</span><span>外来者</span><span>爪牙</span><span>恶魔</span></div>';
    let nums = Object.keys(cfg).map(Number).sort((a,b)=>a-b);
    for(let n of nums) {
        let cc = cfg[n];
        let cls = n === cur ? ' rc-row rc-current' : ' rc-row';
        rows += '<div class="' + cls + '">' +
            '<span>' + n + '人</span>' +
            '<span class="tf">' + cc.townsfolk + '</span>' +
            '<span class="os">' + cc.outsider + '</span>' +
            '<span class="mi">' + cc.minion + '</span>' +
            '<span class="de">' + cc.demon + '</span>' +
            '</div>';
    }
    document.getElementById('rcTable').innerHTML = rows;
}
function stepGame() {
    if(_stepping) return;
    _stepping = true;
    fetch('/step',{method:'POST'}).then(r=>r.json()).then(d=>{ updateBoard(d); _stepping=false; }).catch(()=>{ _stepping=false; });
}
function toggleAuto() {
    let btn = document.getElementById('btnAuto');
    if(autoPlaying) { autoPlaying=false; clearInterval(autoTimer); btn.textContent='▶ 自动播放'; btn.style.background='linear-gradient(135deg,#2196F3,#1a88e8)'; }
    else { autoPlaying=true; btn.textContent='⏸ 暂停'; btn.style.background='linear-gradient(135deg,#ff9800,#f57c00)';
        autoTimer=setInterval(()=>{ if(document.getElementById('btnStep').disabled){toggleAuto();return;} if(_stepping) return; stepGame(); },1600); }
}
function resetGame() { if(autoPlaying) toggleAuto(); document.getElementById('overlay').classList.remove('show'); fetch('/reset',{method:'POST'}).then(r=>r.json()).then(updateBoard); }
function closeOverlay() { document.getElementById('overlay').classList.remove('show'); }
function setPlayerCount(n) { fetch('/set_players/'+n,{method:'POST'}).then(r=>r.json()).then(updateBoard); }
fetch('/state').then(r=>r.json()).then(updateBoard);

/* ========== 日志增强: 类型分类 + 过滤 + 自动滚动 ========== */
let _logFilter = 'all', _autoScroll = true;

function getLogType(msg) {
    if(/====+/.test(msg)) return 'phase';
    if(/天黑|夜幕|闭眼/.test(msg)) return 'night';
    if(/天亮了|黎明|日出|天亮/.test(msg)) return 'day';
    if(/死亡|处决|被杀|毒死/.test(msg)) return 'death';
    if(/\\[第\\d+轮\\]/.test(msg) || /私聊/.test(msg)) return 'chat';
    if(/公聊|公开讨论/.test(msg)) return 'chat';
    if(/提名|投票|表决|得票|票数/.test(msg)) return 'vote';
    if(/使用|发动|技能|查看|得知|选择|首夜信息|爪牙信息|恶魔信息/.test(msg)) return 'action';
    if(/获胜|胜利|结束|阵营|游戏结束/.test(msg)) return 'result';
    if(/设置|角色列表|角色已分配/.test(msg)) return 'system';
    return 'system';
}

function renderLog(log) {
    let container = document.getElementById('logEntries');
    if(!log || log.length === 0) {
        container.innerHTML = '<div class="log-entry entry-system">等待开始...</div>';
        return;
    }
    let h = '';
    for(let e of log) {
        let type = getLogType(e);
        let matchFilter = (_logFilter === 'all' || type === _logFilter);
        let hled = e.replace(/(玩家\\d+)/g, '<span style="color:#4fc3f7;font-weight:bold;">$1</span>');
        hled = hled.replace(/\\[([^\\]]+)\\]/g, '<span style="color:#ffd700;font-weight:bold;">[$1]</span>');
        h += '<div class="log-entry entry-'+type+(matchFilter?'':' hidden')+'">'+hled+'</div>';
    }
    container.innerHTML = h;
    if(_autoScroll) {
        let area = document.getElementById('logArea');
        if(area) area.scrollTop = area.scrollHeight;
    }
}

// 重新计算圆形位置（窗口resize时）
window.addEventListener('resize', function() {
    let container = document.getElementById('circleContainer');
    if(!container || !container.children.length) return;
    let cards = container.children;
    let n = cards.length;
    let cardSize = parseInt(cards[0].style.width) || Math.max(68, Math.min(110, 180 - n * 8));
    let minGap = cardSize + 12;
    let idealR = minGap / (2 * Math.sin(Math.PI / n));
    let maxR = Math.min(container.offsetWidth / 2.3, container.offsetHeight / 2.3, 240);
    let radius = Math.min(idealR, maxR);
    let cx = container.offsetWidth / 2, cy = container.offsetHeight / 2;
    for(let i = 0; i < n; i++) {
        let angle = (i / n) * 2 * Math.PI - Math.PI / 2;
        cards[i].style.left = (cx + radius * Math.cos(angle)) + 'px';
        cards[i].style.top = (cy + radius * Math.sin(angle)) + 'px';
    }
});

function filterLog(type, btn) {
    _logFilter = type;
    document.querySelectorAll('#logFilter button').forEach(b => b.classList.remove('active'));
    if(btn) btn.classList.add('active');
    if(window._fullLogData) renderLog(window._fullLogData);
}

function exportDoc() { window.open('/export'); }

function toggleAutoScroll() {
    _autoScroll = document.getElementById('autoScrollToggle').checked;
}

// 键盘快捷键: Space=下一步, R=重置
document.addEventListener('keydown', function(e) {
    if(e.target.tagName === 'INPUT' || e.target.tagName === 'SELECT') return;
    if(e.key === ' ' || e.key === 'Space') { e.preventDefault(); stepGame(); }
    if(e.key === 'r' || e.key === 'R') { resetGame(); }
});
</script>
</body>
</html>"""

@app.route('/')
def index(): return render_template_string(HTML)

@app.route('/state')
def state(): 
    resp = jsonify(ctrl.get_state())
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate'
    return resp

@app.route('/step', methods=['POST'])
def step():
    try:
        return jsonify(ctrl.step())
    except Exception as e:
        import traceback
        traceback.print_exc()
        try:
            with open(os.path.join(os.path.dirname(__file__), 'crash.log'), 'a', encoding='utf-8') as f:
                f.write(f"\n=== /step ERROR ===\n")
                traceback.print_exc(file=f)
        except: pass
        return jsonify({'error': str(e), 'finished': True, 'phase_name': '⚠️ 游戏出错', 'alive_count': 0, 'players': [], 'log': [], 'chat_threads': []})

@app.route('/reset', methods=['POST'])
def reset():
    global ctrl; ctrl = GameController(ctrl.num_players)
    return jsonify(ctrl.get_state())

@app.route('/set_players/<int:n>', methods=['POST'])
def set_players(n):
    global ctrl; ctrl = GameController(n)
    return jsonify(ctrl.get_state())

@app.route('/storyteller_info')
def storyteller_info():
    g = ctrl.game
    all_agents = g.registry.all_agents()
    team_cn = {'townsfolk': '镇民', 'outsider': '外来者', 'minion': '爪牙', 'demon': '恶魔'}
    info_display_map = {
        'seer': lambda v: f"查验了 {v[0]}，{'发现恶魔' if v[1] else '无恶魔'}",
        'washerwoman': lambda v: f"{v[0]} 和 {v[1]} 中有一人是 {v[2]}",
        'librarian': lambda v: f"{v[0]} 和 {v[1]} 中有一人是 {v[2]}" if v[0] != '无' else "本局没有外来者",
        'investigator': lambda v: f"{v[0]} 和 {v[1]} 中有一人是 {v[2]}",
        'empathy': lambda v: f"左右邻居中有 {v} 个邪恶玩家",
        'chef': lambda v: f"有 {v} 对相邻的邪恶玩家",
        'undertaker': lambda v: f"处决玩家的身份是 {v}",
        'ravenkeeper': lambda v: f"查看了 {v[0]}，身份是 {v[1]}",
    }
    players_info = []
    for a in all_agents:
        info = BOTC_ROLES.get(a.role, {})
        known = a.game_state.get('known_info', {})
        info_items = []
        for k, v in known.items():
            if k in info_display_map:
                info_items.append({'key': k, 'text': info_display_map[k](v)})
            elif k == 'spy_info':
                info_items.append({'key': 'spy_info', 'text': '查看了所有玩家的身份'})
            elif k == 'demon':
                info_items.append({'key': 'demon', 'text': f'恶魔是 {v}'})
            elif k == 'minions':
                info_items.append({'key': 'minions', 'text': f'爪牙是 {", ".join(v)}'})
            elif k == 'master':
                info_items.append({'key': 'master', 'text': f'我的主人是 {v}'})
            elif k == 'fake_roles':
                info_items.append({'key': 'fake_roles', 'text': f'可伪装角色: {", ".join(v)}'})
        # 合并夜间信息历史
        nih = g.game_record.get("night_info_history", {})
        history_items = []
        for day_key in sorted(nih.keys(), key=lambda x: int(x.replace('day_', ''))):
            if a.name in nih[day_key]:
                for entry in nih[day_key][a.name]:
                    day_num = int(day_key.replace('day_', ''))
                    # 首夜信息不显示天数；kill信息显示第几晚
                    if entry['key'] == 'night_kill':
                        text = entry['text']  # 已包含"第X晚"
                    elif day_num == 1:
                        text = entry['text']
                    else:
                        text = f"(第{day_num}天) {entry['text']}"
                    history_items.append({'key': entry['key'], 'text': text})
        # 合并免重
        seen_texts = set()
        merged = []
        for item in info_items + history_items:
            if item['text'] not in seen_texts:
                seen_texts.add(item['text'])
                merged.append(item)
        players_info.append({
            'name': a.name,
            'role': a.role,
            'icon': info.get('icon', '?'),
            'team': info.get('team', ''),
            'team_cn': team_cn.get(info.get('team', ''), ''),
            'alive': a.alive,
            'fake_role': a.game_state.get('fake_role', ''),
            'original_role': a.game_state.get('original_role', ''),  # 红唇继承标注
            'info': merged,
        })
    return jsonify({'players': players_info})

from games.blood_on_clocktower.roles import BOTC_ROLES, BOTC_TEAMS, NIGHT_ORDER_FIRST, NIGHT_ORDER_OTHER

@app.route('/role_data')
def role_data():
    roles_list = []
    for role_name, data in BOTC_ROLES.items():
        roles_list.append({
            'name': role_name,
            'team': data.get('team', ''),
            'ability': data.get('ability', ''),
            'icon': data.get('icon', '?'),
            'first_night': data.get('first_night', False),
            'other_nights': data.get('other_nights', False),
        })
    teams = {}
    for team_key, team_cn in [('townsfolk', '镇民'), ('outsider', '外来者'), ('minion', '爪牙'), ('demon', '恶魔')]:
        teams[team_key] = {'name': team_cn, 'roles': BOTC_TEAMS.get(team_key, [])}
    return jsonify({
        'roles': roles_list,
        'teams': teams,
        'night_order_first': [{'step': s, 'desc': d} for s, d in NIGHT_ORDER_FIRST],
        'night_order_other': [{'step': s, 'desc': d} for s, d in NIGHT_ORDER_OTHER],
    })

@app.route('/export')
def export_doc():
    g = ctrl.game
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    result = g.game_record.get('result', '进行中')
    winner_map = {'townsfolk_win': '善良阵营 (镇民) 胜利', 'evil_win': '邪恶阵营 (恶魔/爪牙) 胜利',
                  'good_win': '善良阵营 胜利', 'demon_win': '恶魔 胜利', 'minion_win': '爪牙 胜利'}
    result_cn = winner_map.get(result, result)
    team_cn = {'townsfolk': '镇民', 'outsider': '外来者', 'minion': '爪牙', 'demon': '恶魔'}
    all_agents = g.registry.all_agents()
    player_names = set(a.name for a in all_agents)
    import re as _re

    # 人名高亮辅助: 公聊=深青+绿名, 私聊=深紫+红名
    def _add_colored_line(paragraph, text, chat_type="public"):
        text_color = RGBColor(0x00, 0x6B, 0x6B) if chat_type == "public" else RGBColor(0x80, 0x00, 0x80)
        name_color = RGBColor(0x00, 0x80, 0x00) if chat_type == "public" else RGBColor(0x8B, 0x00, 0x00)
        m = _re.match(r'^(\s*)(\S+?):\s*(.*)', text)
        if m and m.group(2) in player_names:
            if m.group(1):
                paragraph.add_run(m.group(1))
            rn = paragraph.add_run(m.group(2))
            rn.bold = True
            rn.font.color.rgb = name_color
            paragraph.add_run(': ')
            rc = paragraph.add_run(m.group(3))
            rc.font.color.rgb = text_color
        else:
            paragraph.add_run(text)

    # ═══════════════ 1. 标题 & 概览 ═══════════════
    title = doc.add_heading('血染钟楼 游戏记录', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')
    doc.add_paragraph(f'游戏结果: {result_cn}').bold = True
    doc.add_paragraph(f'剧本: {getattr(g, "script", "暗流涌动")} | 玩家: {ctrl.num_players} 人 | 回合: {g.game_record.get("total_rounds", "?")}')
    doc.add_paragraph('')

    # ═══════════════ 2. 板子配置 & 玩家角色 ═══════════════
    doc.add_heading('一、板子配置与玩家角色', level=1)
    doc.add_paragraph(f'剧本: {getattr(g, "script", "暗流涌动")}')
    doc.add_paragraph('')
    doc.add_heading('玩家角色一览（含首夜信息）', level=2)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'
    for i, txt in enumerate(['玩家', '角色', '阵营', '状态', '原身份/备注', '首夜信息']):
        table.rows[0].cells[i].text = txt
    for a in all_agents:
        row = table.add_row().cells
        row[0].text = a.name
        info = BOTC_ROLES.get(a.role, {})
        row[1].text = f"{info.get('icon', '?')} {a.role}"
        row[2].text = team_cn.get(info.get('team', ''), '')
        row[3].text = '存活' if a.alive else '死亡'
        # 原身份/备注: 优先展示红唇继承，其次酒鬼伪装
        notes = []
        orig = a.game_state.get('original_role', '')
        drunk = a.game_state.get('fake_role', '')
        if orig:
            notes.append(f'原角色: {orig}')
        if drunk:
            notes.append(f'酒鬼(以为{drunk})')
        row[4].text = '；'.join(notes) if notes else '-'
        known = a.game_state.get('known_info', {})
        info_parts = []
        # 仅首夜信息（后续轮次的查验/共情/送葬/守鸦不计入）
        first_night_keys = {'washerwoman','librarian','investigator','chef','seer','empathy','demon','minions','fake_roles','spy_info','spy_shared'}
        info_map = {
            'seer': lambda v: f"查验 {v[0]}，{'发现恶魔' if v[1] else '无恶魔'}",
            'washerwoman': lambda v: f"{v[0]}/{v[1]} 中有一人是 {v[2]}",
            'librarian': lambda v: f"{v[0]}/{v[1]} 中有一人是 {v[2]}" if v[0] != '无' else "无外来者",
            'investigator': lambda v: f"{v[0]}/{v[1]} 中有一人是 {v[2]}",
            'empathy': lambda v: f"邻座有 {v} 个邪恶",
            'chef': lambda v: f"邪恶相邻数: {v}",
            'undertaker': lambda v: f"死者身份: {v}",
            'ravenkeeper': lambda v: f"查看 {v[0]} = {v[1]}",
            'demon': lambda v: f"恶魔: {v}",
            'minions': lambda v: f"爪牙: {', '.join(v)}",
            'master': lambda v: f"主人: {v}",
        }
        for k, v in known.items():
            if k not in first_night_keys:
                continue
            if k in info_map:
                info_parts.append(info_map[k](v))
            elif k == 'fake_roles':
                info_parts.append(f"可伪装: {', '.join(v)}")
        row[5].text = '；'.join(info_parts) if info_parts else '-'
    doc.add_paragraph('')

    # ═══════════════ 3. 按天整理每日流程 ═══════════════
    doc.add_heading('二、每日流程', level=1)
    raw_log = g.storyteller_log or []

    # 按天切分 + 分类
    day_entries = {}
    current_day = None
    for line in raw_log:
        cl = line.strip()
        if not cl:
            continue
        m = _re.search(r'第(\d+)晚\(天黑|第(\d+)天\(白天\)|游戏结束', cl)
        if m:
            d = m.group(1) or m.group(2)
            if d:
                current_day = int(d)
            day_entries.setdefault(current_day, {
                'public_chat': [], 'nominations': [], 'abilities': [],
                'deaths': '', 'night_info_shown': False
            })
            if '游戏结束' in cl:
                day_entries[current_day].setdefault('end_result', []).append(cl)
            continue
        if current_day is None:
            continue
        de = day_entries[current_day]
        # 死亡公告
        if cl.startswith('昨晚死亡') or cl.startswith('平安夜'):
            de['deaths'] = cl
            continue
        # 私聊详情行（带序号）跳过: [1] A 对 B 说: ...
        if _re.match(r'^\[\d+\]', cl):
            continue
        # 公聊发言: "  PlayerName: message"（保留原始缩进用于判断格式）
        if line.startswith('  ') and not cl.startswith('['):
            cm = _re.match(r'^  ([^\[(]+?):\s*(.*)', line)
            if cm:
                nm = cm.group(1).strip()
                if nm in player_names:
                    if '[提名发言]' not in cl and '[辩护发言]' not in cl:
                        de['public_chat'].append(cl)
                        continue
        # 提名相关
        if _re.search(r'提名了|\[提名发言\]|\[辩护发言\]|\[投票\]|\[计票\]|被处决|未被处决|处决者:', cl):
            de['nominations'].append(cl)
            continue
        # 角色技能标记行
        if cl.startswith('[') and ']' in cl[:20]:
            de['abilities'].append(cl)
            continue

    # ---- 输出每日 ----
    for day_num in sorted(day_entries.keys()):
        de = day_entries[day_num]
        doc.add_heading(f'■ 第 {day_num} 天', level=2)

        # 夜晚小结
        night_info_shown = False
        day_key = f"day_{day_num}"
        nih = g.game_record.get("night_info_history", {}).get(day_key, {})
        deaths = de.get('deaths', '')
        if nih or deaths:
            doc.add_heading(f'🌙 夜晚', level=3)
            if deaths:
                doc.add_paragraph(deaths)
            if nih:
                for pname in sorted(nih.keys()):
                    for entry in nih[pname]:
                        doc.add_paragraph(f'{pname}: {entry["text"]}')
            night_info_shown = True

        # 公聊
        if de['public_chat']:
            doc.add_heading(f'💬 公聊', level=3)
            for sp in de['public_chat']:
                p = doc.add_paragraph()
                _add_colored_line(p, sp, "public")

        # 提名
        if de['nominations']:
            doc.add_heading(f'🗳️ 提名与投票', level=3)
            # 从 structured data 取提名详情
            day_votes = g.game_record.get("vote_history", {}).get(day_key, [])
            day_noms = g.game_record.get("nomination_history", {}).get(day_key, [])
            # 先输出结构化提名（含完整投票明细）
            for nom in day_noms:
                ntr = nom.get('target', '?')
                nmr = nom.get('nominator', '?')
                rs = nom.get('result', '')
                rs_icon = '🔴 被处决' if rs == 'executed' else '🟢 未被处决'
                p = doc.add_paragraph()
                run = p.add_run(f'{nmr} 提名 {ntr} → {rs_icon}')
                run.bold = True
                if nom.get('nominator_speech'):
                    doc.add_paragraph(f'  [提名发言] {nmr}: {nom["nominator_speech"]}')
                if nom.get('defense_speech'):
                    doc.add_paragraph(f'  [辩护发言] {ntr}: {nom["defense_speech"]}')
                related_votes = [v for v in day_votes if v.get('target') == ntr]
                if related_votes:
                    doc.add_paragraph(f'  投票 ({len(related_votes)} 票):')
                    for v in related_votes:
                        doc.add_paragraph(f'    {v.get("voter", "?")} 投 {v.get("target", "?")}')
                doc.add_paragraph('')
            # 补充 raw_log 中未出现在结构化提名的行（如计票结果）
            structured_nom_keys = set()
            for nom in day_noms:
                structured_nom_keys.add(nom.get('nominator', '') + '提名了' + nom.get('target', ''))
            for nline in de['nominations']:
                # 如果该提名已在结构化数据中显示，跳过
                extra = True
                for nom in day_noms:
                    if nom.get('nominator', '') + '提名了' + nom.get('target', '') in nline:
                        extra = False
                        break
                if extra:
                    p = doc.add_paragraph()
                    run = p.add_run(nline)
                    run.italic = True

        # 角色技能
        if de['abilities']:
            doc.add_heading(f'⚔️ 角色技能使用', level=3)
            for ab in de['abilities']:
                doc.add_paragraph(ab)

        # 游戏结束
        if 'end_result' in de:
            for er in de['end_result']:
                doc.add_paragraph(er)

    # ═══════════════ 4. 私聊记录 ═══════════════
    pch = g.game_record.get('private_chat_history', {})
    if pch:
        doc.add_heading('三、私聊记录', level=1)
        doc.add_paragraph('（以下按对话双方分组，每条消息格式为 "说话人: 内容"）')
        for thread_key in sorted(pch.keys()):
            msgs = pch[thread_key]
            display = thread_key.replace('\U0001f504', ' ↔ ')
            p = doc.add_paragraph()
            run = p.add_run(f'【{display}】')
            run.bold = True
            for msg in msgs:
                p = doc.add_paragraph()
                _add_colored_line(p, f'  {msg.get("speaker", "?")}: {msg.get("text", "")}', "private")

    # ═══════════════ 5. 游戏结果 ═══════════════
    doc.add_page_break()
    doc.add_heading('四、游戏结果', level=1)
    doc.add_paragraph(f'结果: {result_cn}')
    doc.add_paragraph(f'总回合: {g.game_record.get("total_rounds", "?")}')
    survivors = [a.name for a in all_agents if a.alive]
    doc.add_paragraph(f'幸存者: {", ".join(survivors) if survivors else "无 (全灭)"}')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    return send_file(buf, as_attachment=True, download_name=f'血染钟楼_游戏记录_{ts}.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/private_chat')
def private_chat():
    resp = make_response(render_template_string(PRIVATE_CHAT_HTML))
    resp.headers['Cache-Control'] = 'no-store'
    return resp

PRIVATE_CHAT_HTML = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>血染钟楼 - 说书人面板</title>
<style>
* { margin:0; padding:0; box-sizing:border-box; }
body { font-family:'Microsoft YaHei',sans-serif; background:#0a0a1a; color:#e0e0e0; height:100vh; overflow:hidden; display:flex; flex-direction:column; }
.tab-bar { background:linear-gradient(135deg,#1a1a3e,#0d0d2a); display:flex; border-bottom:2px solid #ffd70033; flex-shrink:0; }
.tab-bar .tab { padding:12px 24px; font-size:14px; font-weight:bold; cursor:pointer; color:#888; border-bottom:3px solid transparent; transition:all .15s; user-select:none; }
.tab-bar .tab:hover { color:#ccc; background:#1a1a3e; }
.tab-bar .tab.active { color:#ffd700; border-bottom-color:#ffd700; background:rgba(255,215,0,.05); }
.tab-content { flex:1; display:flex; flex-direction:column; min-height:0; }
.tab-pane { display:none; flex:1; flex-direction:column; min-height:0; }
.tab-pane.active { display:flex; }
/* ========== 私聊记录 (Tab 1) ========== */
.chat-layout { display:grid; grid-template-columns:320px 1fr; gap:0; flex:1; min-height:0; overflow:hidden; }
.thread-panel { background:#0d0d1a; border-right:1px solid #222; display:flex; flex-direction:column; overflow:hidden; }
.thread-panel .panel-title { color:#ffd700; font-size:13px; padding:12px 14px; border-bottom:1px solid #1a1a2e; flex-shrink:0; font-weight:bold; }
.thread-list { flex:1; overflow-y:auto; }
.thread-list::-webkit-scrollbar { width:5px; }
.thread-list::-webkit-scrollbar-track { background:transparent; }
.thread-list::-webkit-scrollbar-thumb { background:#333; border-radius:3px; }
.thread-item { display:flex; align-items:center; gap:10px; padding:10px 14px; cursor:pointer; transition:background .12s; border-bottom:1px solid #1a1a2e; }
.thread-item:hover { background:#1a1a3e; }
.thread-item.active { background:#1a2a1a; border-left:3px solid #4CAF50; }
.thread-item .t-avatar { width:32px; height:32px; border-radius:50%; display:flex; align-items:center; justify-content:center; font-size:14px; font-weight:bold; color:#fff; flex-shrink:0; }
.thread-item .t-info { flex:1; min-width:0; }
.thread-item .t-name { font-size:13px; font-weight:bold; color:#e0e0e0; }
.thread-item .t-preview { font-size:11px; color:#888; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; margin-top:2px; }
.thread-item .t-right { text-align:right; flex-shrink:0; }
.thread-item .t-day { font-size:10px; color:#555; }
.thread-item .unread-badge { display:inline-block; background:#f44336; color:#fff; font-size:9px; padding:1px 5px; border-radius:8px; font-weight:bold; margin-top:3px; }
.chat-panel { display:flex; flex-direction:column; overflow:hidden; background:#0d0d1a; }
.chat-header { background:linear-gradient(135deg,#1a1a3e,#0d0d2a); padding:12px 16px; display:flex; align-items:center; gap:8px; border-bottom:1px solid #222; flex-shrink:0; }
.chat-header .chat-title { color:#fff; font-size:15px; font-weight:bold; flex:1; }
.chat-header .chat-empty-hint { color:#666; font-size:12px; }
.chat-bubbles { padding:14px 16px; overflow-y:auto; flex:1; display:flex; flex-direction:column; gap:10px; }
.chat-bubbles::-webkit-scrollbar { width:5px; }
.chat-bubbles::-webkit-scrollbar-track { background:transparent; }
.chat-bubbles::-webkit-scrollbar-thumb { background:#333; border-radius:3px; }
.chat-bubble-row { display:flex; align-items:flex-start; gap:10px; max-width:80%; }
.chat-bubble-row.right { flex-direction:row-reverse; align-self:flex-end; }
.chat-bubble-row.left { align-self:flex-start; }
.chat-bubble-row .b-avatar { width:30px; height:30px; border-radius:50%; display:flex; align-items:center; justify-content:center; font-size:13px; font-weight:bold; color:#fff; flex-shrink:0; }
.chat-bubble-row .b-body { display:flex; flex-direction:column; }
.chat-bubble-row .b-speaker { font-size:11px; color:#888; margin-bottom:2px; padding-left:4px; }
.chat-bubble-row.right .b-speaker { text-align:right; padding-right:4px; }
.chat-bubble-row .bubble { padding:8px 14px; font-size:13px; line-height:1.5; word-break:break-word; position:relative; }
.chat-bubble-row.left .bubble { background:#1a1a3e; color:#e0e0e0; border-radius:10px 10px 10px 4px; border:1px solid #333; }
.chat-bubble-row.right .bubble { background:#1a4a1a; color:#e0e0e0; border-radius:10px 10px 4px 10px; border:1px solid #2a6a2a; }
.chat-bubble-row .b-time { font-size:10px; color:#555; margin-top:3px; padding-left:4px; }
.chat-bubble-row.right .b-time { text-align:right; padding-right:4px; }
.chat-empty { text-align:center; color:#555; padding:40px 20px; font-size:13px; }
.chat-footer { padding:8px 16px; border-top:1px solid #222; flex-shrink:0; text-align:center; }
.chat-footer a { color:#ffd700; text-decoration:none; font-size:13px; }
.chat-footer a:hover { text-decoration:underline; }
/* ========== 说书人信息 (Tab 2) ========== */
.st-container { flex:1; overflow-y:auto; padding:20px; }
.st-container::-webkit-scrollbar { width:6px; }
.st-container::-webkit-scrollbar-track { background:transparent; }
.st-container::-webkit-scrollbar-thumb { background:#333; border-radius:3px; }
.st-grid { display:grid; grid-template-columns:repeat(auto-fill,minmax(320px,1fr)); gap:14px; }
.st-card { background:linear-gradient(135deg,#111125,#0d0d1a); border:1px solid #2a2a4e; border-radius:12px; overflow:hidden; }
.st-card-header { padding:12px 16px; display:flex; align-items:center; gap:10px; border-bottom:1px solid #1a1a2e; }
.st-card-header .st-icon { font-size:24px; }
.st-card-header .st-name { font-size:16px; font-weight:bold; color:#e0e0e0; }
.st-card-header .st-role { font-size:12px; padding:2px 10px; border-radius:10px; font-weight:bold; }
.st-card-header .st-role.townsfolk { background:#4a90d944; color:#7ab8ff; }
.st-card-header .st-role.outsider { background:#7b68ee44; color:#a89bff; }
.st-card-header .st-role.minion { background:#d94a4a44; color:#ff7a7a; }
.st-card-header .st-role.demon { background:#ff444444; color:#ff8888; }
.st-card-header .st-tag { font-size:11px; padding:2px 8px; border-radius:8px; margin-left:auto; }
.st-card-header .st-tag.alive { background:#4CAF5022; color:#8BC34A; }
.st-card-header .st-tag.dead { background:#66666622; color:#999; }
.st-card-body { padding:12px 16px; }
.st-info-item { padding:6px 0; border-bottom:1px solid #1a1a2e; font-size:13px; color:#ccc; display:flex; align-items:flex-start; gap:8px; }
.st-info-item:last-child { border-bottom:none; }
.st-info-item .st-info-icon { color:#ffd700; flex-shrink:0; }
.st-info-empty { color:#555; font-size:12px; padding:8px 0; text-align:center; }
.st-fake-role { font-size:12px; color:#ffa940; padding:6px 0 0; }
::-webkit-scrollbar { width:6px; }
::-webkit-scrollbar-track { background:transparent; }
::-webkit-scrollbar-thumb { background:#333; border-radius:3px; }
@media (max-width:700px) {
    .chat-layout { grid-template-columns:1fr; }
    .thread-panel { max-height:180px; border-right:none; border-bottom:1px solid #222; }
    .st-grid { grid-template-columns:1fr; }
}
/* ========== 角色图鉴表格 ========== */
.roles-table { width:100%; border-collapse:collapse; font-size:13px; }
.roles-table th { padding:8px 12px; text-align:left; color:#ffd700; border-bottom:2px solid #ffd70044; background:#1a1a3e; }
.roles-table td { padding:6px 12px; border-bottom:1px solid #1a1a2e; color:#aaa; }
.roles-table tr:hover td { background:#1a1a3e44; }
.roles-table .rname { color:#e0e0e0; font-weight:bold; }
</style>
</head>
<body>
<div class="tab-bar">
    <div class="tab active" onclick="switchTab('chat',this)">💬 私聊记录</div>
    <div class="tab" onclick="switchTab('storyteller',this)">📜 说书人信息</div>
    <div class="tab" onclick="switchTab('nightorder',this)">🌙 夜晚顺序</div>
    <div class="tab" onclick="switchTab('roleabilities',this)">📋 角色能力</div>
    <div class="tab" onclick="switchTab('nomination',this)">🗳️ 提名记录</div>
</div>
<div class="tab-content">
    <!-- Tab 1: 私聊记录 -->
    <div class="tab-pane active" id="tabChat">
        <div class="chat-layout">
            <div class="thread-panel">
                <div class="panel-title">📋 会话列表</div>
                <div id="threadList" class="thread-list"></div>
            </div>
            <div class="chat-panel">
                <div class="chat-header">
                    <span class="chat-title" id="chatTitle">选择一条会话</span>
                    <span class="chat-empty-hint" id="chatHint">👈 左侧选择</span>
                </div>
                <div id="chatBubbles" class="chat-bubbles"><div class="chat-empty">请在左侧选择一条私聊记录</div></div>
                <div class="chat-footer" id="chatFooter" style="display:none">
                    <a href="#" onclick="scrollToThreads();return false;">← 返回会话列表</a>
                </div>
            </div>
        </div>
    </div>
    <!-- Tab 2: 说书人信息 -->
    <div class="tab-pane" id="tabStoryteller">
        <div class="st-container" id="stContainer">
            <div class="chat-empty" style="padding:60px 20px;">加载中...</div>
        </div>
    </div>
    <!-- Tab 3: 夜晚顺序 -->
    <div class="tab-pane" id="tabNightorder">
        <div class="st-container" id="nightOrderContainer">
            <div class="chat-empty" style="padding:60px 20px;">加载中...</div>
        </div>
    </div>
    <!-- Tab 4: 角色能力 -->
    <div class="tab-pane" id="tabRoleabilities">
        <div class="st-container" id="roleAbilitiesContainer">
            <div class="chat-empty" style="padding:60px 20px;">加载中...</div>
        </div>
    </div>
    <!-- Tab 5: 提名记录 -->
    <div class="tab-pane" id="tabNomination">
        <div class="st-container" id="nominationContainer">
            <div class="chat-empty" style="padding:60px 20px;">加载中...</div>
        </div>
    </div>
</div>
<script>
let chatData = null;
let currentThreadKey = null;

function switchTab(tab, btn) {
    document.querySelectorAll('.tab-bar .tab').forEach(t => t.classList.remove('active'));
    document.querySelectorAll('.tab-pane').forEach(p => p.classList.remove('active'));
    if(btn) btn.classList.add('active');
    document.getElementById('tab' + tab.charAt(0).toUpperCase() + tab.slice(1)).classList.add('active');
    if(tab === 'chat') loadChat();
    else if(tab === 'storyteller') loadStorytellerInfo();
    else if(tab === 'nightorder') loadNightOrder();
    else if(tab === 'roleabilities') loadRoleAbilities();
    else if(tab === 'nomination') loadNominationHistory();
}

function loadChat() {
    fetch('/state').then(r => r.json()).then(d => {
        chatData = d;
        renderThreads(d.chat_threads || []);
        if(currentThreadKey) {
            let match = (d.chat_threads || []).find(t => t.key === currentThreadKey);
            if(match) renderBubbles(match);
        }
    }).catch(e => {
        document.getElementById('threadList').innerHTML = '<div class="chat-empty">加载失败: ' + e.message + '</div>';
    });
}

function renderThreads(threads) {
    let list = document.getElementById('threadList');
    if(!threads || threads.length === 0) {
        list.innerHTML = '<div class="chat-empty">暂无记录</div>';
        return;
    }
    let cols = ['#07c160','#1485ee','#ff6b6b','#ffa940','#ab47bc','#26c6da','#8d6e63','#78909c','#5c6bc0','#ec407a','#66bb6a'];
    let h = '';
    for(let t of threads) {
        let msgs = t.messages || [];
        let speakers = [...new Set(msgs.map(m => m.speaker))];
        let label = t.pair || speakers.join(' ↔ ') || '?';
        let avatar = (speakers[0] || label).charAt(0);
        let color = cols[threads.indexOf(t) % cols.length];
        let preview = msgs.length > 0 ? (msgs[msgs.length-1].text || '') : '';
        if(preview.length > 30) preview = preview.substring(0,30) + '…';
        let active = (currentThreadKey === t.key) ? ' active' : '';
        h += '<div class="thread-item' + active + '" data-key="' + t.key.replace(/"/g, '&quot;') + '">' +
            '<div class="t-avatar" style="background:' + color + '">' + escapeHtml(avatar) + '</div>' +
            '<div class="t-info"><div class="t-name">' + escapeHtml(label) + '</div><div class="t-preview">' + escapeHtml(preview) + '</div></div>' +
            '<div class="t-right"><div class="t-day">' + escapeHtml(t.day || '') + '</div></div></div>';
    }
    list.innerHTML = h;
}

function openThread(key) {
    currentThreadKey = key;
    document.querySelectorAll('.thread-item').forEach(el => el.classList.remove('active'));
    let el = document.querySelector('.thread-item[data-key="' + key.replace(/"/g, '') + '"]');
    if(el) el.classList.add('active');
    document.getElementById('chatFooter').style.display = 'block';
    if(window.innerWidth <= 700) {
        setTimeout(() => document.querySelector('.chat-panel').scrollIntoView({ behavior:'smooth' }), 100);
    }
    if(chatData) {
        let match = chatData.chat_threads.find(t => t.key === key);
        if(match) renderBubbles(match);
    }
}

function renderBubbles(thread) {
    let container = document.getElementById('chatBubbles');
    document.getElementById('chatTitle').textContent = thread.pair || [...new Set((thread.messages||[]).map(m=>m.speaker))].join(' ↔ ') || '对话';
    document.getElementById('chatHint').textContent = '';
    if(!thread.messages || thread.messages.length === 0) {
        container.innerHTML = '<div class="chat-empty">暂无消息</div>';
        return;
    }
    let speakers = [...new Set(thread.messages.map(m => m.speaker))];
    let leftS = speakers[0], rightS = speakers.length > 1 ? speakers[1] : speakers[0];
    let cols = ['#07c160','#1485ee','#ff6b6b','#ffa940','#ab47bc','#26c6da','#8d6e63','#78909c','#5c6bc0','#ec407a','#66bb6a'];
    let used = {};
    let h = '';
    for(let m of thread.messages) {
        let isRight = m.speaker === rightS && rightS !== leftS;
        if(!used[m.speaker]) used[m.speaker] = cols[Object.keys(used).length % cols.length];
        let now = new Date();
        let ts = String(now.getHours()).padStart(2,'0') + ':' + String(now.getMinutes()).padStart(2,'0') + ':' + String(now.getSeconds()).padStart(2,'0');
        h += '<div class="chat-bubble-row ' + (isRight ? 'right' : 'left') + '">' +
            '<div class="b-avatar" style="background:' + used[m.speaker] + '">' + escapeHtml(m.speaker.charAt(0)) + '</div>' +
            '<div class="b-body"><div class="b-speaker">' + escapeHtml(m.speaker) + '</div>' +
            '<div class="bubble">' + escapeHtml(m.text) + '</div>' +
            '<div class="b-time">' + ts + '</div></div></div>';
    }
    container.innerHTML = h;
    container.scrollTop = container.scrollHeight;
}

function escapeHtml(s) {
    if(!s) return '';
    return s.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;');
}

function loadStorytellerInfo() {
    let container = document.getElementById('stContainer');
    fetch('/storyteller_info').then(r=>r.json()).then(d => {
        if(!d.players || d.players.length === 0) {
            container.innerHTML = '<div class="chat-empty" style="padding:60px 20px;">暂无信息</div>';
            return;
        }
        let h = '<div class="st-grid">';
        for(let p of d.players) {
            let icon = p.icon || '?';
            let team = p.team || '';
            let status = p.alive ? 'alive' : 'dead';
            let fakeLine = p.fake_role ? '<div class="st-fake-role">🍺 酒鬼，误认为自己：' + escapeHtml(p.fake_role) + '</div>' : '';
            let infoHtml = '';
            if(p.info && p.info.length > 0) {
                for(let item of p.info) {
                    infoHtml += '<div class="st-info-item"><span class="st-info-icon">📌</span><span>' + escapeHtml(item.text) + '</span></div>';
                }
            } else {
                infoHtml = '<div class="st-info-empty">该角色无首夜信息</div>';
            }
            h += '<div class="st-card">' +
                '<div class="st-card-header">' +
                '<span class="st-icon">' + icon + '</span>' +
                '<span class="st-name">' + escapeHtml(p.name) + '</span>' +
                '<span class="st-role ' + team + '">' + escapeHtml(p.role) + '</span>' +
                '<span class="st-tag ' + status + '">' + (p.alive ? '存活' : '死亡') + '</span>' +
                '</div>' +
                '<div class="st-card-body">' + infoHtml + fakeLine + '</div>' +
                '</div>';
        }
        h += '</div>';
        container.innerHTML = h;
    }).catch(() => {
        container.innerHTML = '<div class="chat-empty" style="padding:60px 20px;">加载失败，请重试</div>';
    });
}

function scrollToThreads() {
    document.querySelector('.thread-panel').scrollIntoView({ behavior:'smooth' });
}

function loadNightOrder() {
    let container = document.getElementById('nightOrderContainer');
    Promise.all([
        fetch('/role_data').then(r=>r.json()),
        fetch('/state').then(r=>r.json())
    ]).then(([rd, st]) => {
        let rolePlayers = {};
        for(let p of st.players) {
            let r = p.role;
            if(!rolePlayers[r]) rolePlayers[r] = [];
            rolePlayers[r].push(p.name);
        }
        let stepAliases = {'爪牙信息':'调查员','恶魔信息':'占卜师'};
        let h = '<div style="padding:20px;max-width:750px;">';
        h += '<h3 style="color:#ffd700;font-size:16px;margin-bottom:12px;">🌙 首夜顺序</h3>' +
            '<table style="width:100%;border-collapse:collapse;font-size:13px;margin-bottom:24px;"><thead>' +
            '<tr style="background:#1a1a3e;"><th style="padding:8px 12px;text-align:left;color:#ffd700;border-bottom:2px solid #ffd70044;width:140px;">步骤</th>' +
            '<th style="padding:8px 12px;text-align:left;color:#ffd700;border-bottom:2px solid #ffd70044;">说明</th></tr></thead><tbody>';
        for(let n of rd.night_order_first) {
            let alias = stepAliases[n.step];
            let extra = alias && rolePlayers[alias] ? rolePlayers[alias] : [];
            let matched = rd.roles.find(r => r.name === n.step);
            let pList = (matched && rolePlayers[matched.name] || []).concat(extra);
            let pTag = pList.length > 0 ? ' <span style="color:#8BC34A;font-weight:normal;">(' + pList.join(', ') + ')</span>' : '';
            h += '<tr style="border-bottom:1px solid #1a1a2e;"><td style="padding:6px 12px;color:#e0e0e0;font-weight:bold;">' + escapeHtml(n.step) + pTag + '</td>' +
                '<td style="padding:6px 12px;color:#aaa;">' + escapeHtml(n.desc) + '</td></tr>';
        }
        h += '</tbody></table>';
        h += '<h3 style="color:#ffd700;font-size:16px;margin-bottom:12px;">🌙 其他夜晚顺序</h3>' +
            '<table style="width:100%;border-collapse:collapse;font-size:13px;"><thead>' +
            '<tr style="background:#1a1a3e;"><th style="padding:8px 12px;text-align:left;color:#ffd700;border-bottom:2px solid #ffd70044;width:140px;">步骤</th>' +
            '<th style="padding:8px 12px;text-align:left;color:#ffd700;border-bottom:2px solid #ffd70044;">说明</th></tr></thead><tbody>';
        for(let n of rd.night_order_other) {
            let matched = rd.roles.find(r => r.name === n.step);
            let pList = matched && rolePlayers[matched.name] || [];
            let pTag = pList.length > 0 ? ' <span style="color:#8BC34A;font-weight:normal;">(' + pList.join(', ') + ')</span>' : '';
            h += '<tr style="border-bottom:1px solid #1a1a2e;"><td style="padding:6px 12px;color:#e0e0e0;font-weight:bold;">' + escapeHtml(n.step) + pTag + '</td>' +
                '<td style="padding:6px 12px;color:#aaa;">' + escapeHtml(n.desc) + '</td></tr>';
        }
        h += '</tbody></table></div>';
        container.innerHTML = h;
    }).catch(() => {
        container.innerHTML = '<div class="chat-empty" style="padding:60px 20px;">加载失败</div>';
    });
}

function loadRoleAbilities() {
    let container = document.getElementById('roleAbilitiesContainer');
    Promise.all([
        fetch('/role_data').then(r=>r.json()),
        fetch('/state').then(r=>r.json())
    ]).then(([rd, st]) => {
        let rolePlayers = {};
        for(let p of st.players) {
            let r = p.role;
            if(!rolePlayers[r]) rolePlayers[r] = [];
            rolePlayers[r].push(p.name);
        }
        let h = '<div style="padding:20px;">';
        h += '<h3 style="color:#ffd700;font-size:16px;margin-bottom:12px;">📋 角色能力表</h3>';
        let teamOrder = ['townsfolk','outsider','minion','demon'];
        let teamNames = {'townsfolk':'镇民','outsider':'外来者','minion':'爪牙','demon':'恶魔'};
        let teamColors = {'townsfolk':'#7ab8ff','outsider':'#a89bff','minion':'#ff7a7a','demon':'#ff8888'};
        for(let tk of teamOrder) {
            let roles = rd.roles.filter(r => r.team === tk);
            if(roles.length === 0) continue;
            h += '<h4 style="color:' + teamColors[tk] + ';font-size:14px;margin:14px 0 8px;border-left:3px solid ' + teamColors[tk] + ';padding-left:10px;">' + teamNames[tk] + ' (' + roles.length + '人)</h4>';
            for(let r of roles) {
                let pList = rolePlayers[r.name] || [];
                let pTag = pList.length > 0 ? ' <span style="color:#8BC34A;font-weight:normal;font-size:12px;">(' + pList.join(', ') + ')</span>' : '';
                let nightTag = '';
                if(r.first_night && r.other_nights) nightTag = '<span style="font-size:11px;color:#4CAF50;margin-left:8px;">[首夜+每夜]</span>';
                else if(r.first_night) nightTag = '<span style="font-size:11px;color:#ffa940;margin-left:8px;">[仅首夜]</span>';
                else if(r.other_nights) nightTag = '<span style="font-size:11px;color:#4fc3f7;margin-left:8px;">[每夜]</span>';
                h += '<div style="background:#111125;border:1px solid #1a1a2e;border-radius:8px;padding:10px 14px;margin-bottom:6px;">' +
                    '<div style="display:flex;align-items:center;gap:8px;margin-bottom:4px;">' +
                    '<span>' + r.icon + '</span>' +
                    '<span style="font-weight:bold;color:#e0e0e0;font-size:14px;">' + escapeHtml(r.name) + '</span>' +
                    pTag + nightTag + '</div>' +
                    '<div style="font-size:12px;color:#aaa;line-height:1.5;">' + escapeHtml(r.ability) + '</div></div>';
            }
        }
        h += '</div>';
        container.innerHTML = h;
    }).catch(() => {
        container.innerHTML = '<div class="chat-empty" style="padding:60px 20px;">加载失败</div>';
    });
}

function loadNominationHistory() {
    let container = document.getElementById('nominationContainer');
    fetch('/state').then(r=>r.json()).then(d => {
        let nh = d.nomination_history || {};
        let vh = d.vote_history || {};
        let days = Object.keys(nh).sort();
        if(days.length === 0) {
            container.innerHTML = '<div class="chat-empty" style="padding:60px 20px;">暂无提名记录</div>';
            return;
        }
        let h = '<div style="padding:20px;">';
        for(let dk of days) {
            let dayLabel = dk.replace('day_','第') + '天';
            let noms = nh[dk] || [];
            let dayVotes = vh[dk] || [];
            h += '<h3 style="color:#ffd700;font-size:15px;margin:0 0 8px 0;">' + dayLabel + '</h3>';
            h += '<table style="width:100%;border-collapse:collapse;font-size:13px;margin-bottom:20px;"><thead>' +
                '<tr style="background:#1a1a3e;"><th style="padding:8px 12px;text-align:left;color:#ffd700;border-bottom:2px solid #ffd70044;">提名者</th>' +
                '<th style="padding:8px 12px;text-align:left;color:#ffd700;border-bottom:2px solid #ffd70044;">被提名者</th>' +
                '<th style="padding:8px 12px;text-align:left;color:#ffd700;border-bottom:2px solid #ffd70044;">票数</th>' +
                '<th style="padding:8px 12px;text-align:left;color:#ffd700;border-bottom:2px solid #ffd70044;">投票者</th>' +
                '<th style="padding:8px 12px;text-align:left;color:#ffd700;border-bottom:2px solid #ffd70044;">结果</th></tr></thead><tbody>';
            for(let n of noms) {
                let relatedVotes = dayVotes.filter(v => v.target === n.target);
                let voteCount = relatedVotes.length;
                let voters = relatedVotes.map(v => v.voter).join(', ');
                let resultText = n.result === 'executed' ? '🔴 被处决' : '🟢 未被处决';
                let rowColor = n.result === 'executed' ? 'background:#4a1a1a44;' : '';
                h += '<tr style="border-bottom:1px solid #1a1a2e;' + rowColor + '">' +
                    '<td style="padding:8px 12px;color:#e0e0e0;">' + escapeHtml(n.nominator) + '</td>' +
                    '<td style="padding:8px 12px;color:#e0e0e0;">' + escapeHtml(n.target) + '</td>' +
                    '<td style="padding:8px 12px;color:#ccc;">' + voteCount + '</td>' +
                    '<td style="padding:8px 12px;color:#ccc;">' + escapeHtml(voters || '-') + '</td>' +
                    '<td style="padding:8px 12px;">' + resultText + '</td></tr>';
            }
            h += '</tbody></table>';
        }
        h += '</div>';
        container.innerHTML = h;
    }).catch(() => {
        container.innerHTML = '<div class="chat-empty" style="padding:60px 20px;">加载失败</div>';
    });
}

// 事件委托: 点击会话项
document.getElementById('threadList').addEventListener('click', function(e) {
    var el = e.target;
    while(el && el !== this) {
        if(el.classList && el.classList.contains('thread-item')) {
            openThread(el.dataset.key);
            return;
        }
        el = el.parentNode;
    }
});

// 初始化: 加载默认tab
loadChat();
// 自动刷新
setInterval(function() {
    let activeTab = document.querySelector('.tab-pane.active');
    if(!activeTab) return;
    let id = activeTab.id;
    if(id === 'tabChat') loadChat();
    else if(id === 'tabNomination') loadNominationHistory();
}, 3000);
</script>
</body>
</html>"""

if __name__ == '__main__':
    print('='*50)
    print('  血染钟楼 · 私聊对话框版')
    print('  http://127.0.0.1:5000')
    print('='*50)
    app.run(host='127.0.0.1', port=5000, debug=False, threaded=True)
